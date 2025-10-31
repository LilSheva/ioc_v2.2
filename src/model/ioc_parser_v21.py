"""
Модуль для извлечения IOC из документов V2.1.
ИСПРАВЛЕНИЯ:
- Правильный порядок извлечения хешей (SHA256 → SHA1 → MD5)
- Фильтрация IOC из URI (IP, домены, файлы)
- Фильтрация ложных File
- Email whitelist
"""

import re
from typing import List, Dict, Set, Any, Tuple
from docx import Document


class IOCParser:
    """Парсер для извлечения IOC из .docx файлов V2.1."""
    
    # Whitelist для email (служебные адреса)
    EMAIL_WHITELIST = {
        'otd93@fstec.ru',
        'fstec.ru'  # Домен тоже блокируем
    }
    
    # Blacklist для File (общие слова из угловых скобок)
    FILE_BLACKLIST = {
        '@', 'е', 'i', 'a', 'o', 'u',
        'exe', 'lnk', 'rar', 'zip', 'pdf', 'txt', 'doc', 'docx',
        'AnyDesk', 'TeamViewer', 'UltraVNC',
        'Dr.Web', 'Kaspersky', 'Защита', 'защита',
        'белым', 'черным', 'бэкдор', 'троян', 'червь', 'стилер', 'загрузчик',
        'Панель управления', 'Учетные записи', 'Управление',
        'Базовая защита', 'Компоненты защиты',
        'О дополнительных мерах'
    }
    
    def __init__(self, ioc_config: List[Dict[str, Any]]):
        """Инициализация парсера."""
        self.ioc_config = ioc_config
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлекает весь текст из .docx файла (параграфы и таблицы)."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Извлечение текста из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка при чтении файла {file_path}: {str(e)}")
    
    def extract_from_files(self, file_paths: List[str]) -> str:
        """Извлекает и объединяет текст из нескольких файлов."""
        combined_text = []
        
        for file_path in file_paths:
            text = self.extract_text_from_docx(file_path)
            combined_text.append(text)
        
        return '\n\n'.join(combined_text)
    
    @staticmethod
    def clean_ioc(ioc: str, ioc_type: str) -> str:
        """
        Очищает IOC от обфускации.
        
        Args:
            ioc: Исходный IOC
            ioc_type: Тип IOC
            
        Returns:
            Очищенный IOC
        """
        cleaned = ioc
        
        # Очистка [...] и [:]
        cleaned = cleaned.replace('[.]', '.')
        cleaned = cleaned.replace('[:]', ':')
        cleaned = cleaned.replace('[', '')
        cleaned = cleaned.replace(']', '')
        
        # Дополнительная очистка для URI
        if ioc_type == 'URI':
            # Убираем hxxp -> http, hxxps -> https
            cleaned = cleaned.replace('hxxp://', 'http://')
            cleaned = cleaned.replace('hxxps://', 'https://')
            cleaned = cleaned.replace('hxxp:', 'http:')
            cleaned = cleaned.replace('hxxps:', 'https:')
        
        return cleaned
    
    def extract_components_from_uri(self, uri: str) -> Dict[str, List[str]]:
        """
        Извлекает компоненты из URI для фильтрации.
        
        Returns:
            {'ips': [...], 'domains': [...], 'files': [...]}
        """
        components = {'ips': [], 'domains': [], 'files': []}
        
        # Паттерн для извлечения IP из URI
        ip_pattern = r'(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)(?:\[\\.\]|\\.)){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)'
        
        # Паттерн для извлечения домена из URI
        domain_pattern = r'://([a-zA-Z0-9](?:[a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?(?:\[\\.\]|\\.)+[a-zA-Z]{2,})'
        
        # Паттерн для извлечения имени файла из URI (после последнего /)
        file_pattern = r'/([^/\s?]+\[\\.\](?:rar|zip|exe|pdf|doc|docx|xls|xlsx))(?:[?\s]|$)'
        
        # Извлечение IP
        for match in re.finditer(ip_pattern, uri):
            components['ips'].append(match.group())
        
        # Извлечение домена
        for match in re.finditer(domain_pattern, uri):
            components['domains'].append(match.group(1))
        
        # Извлечение файлов
        for match in re.finditer(file_pattern, uri, re.IGNORECASE):
            components['files'].append(match.group(1))
        
        return components
    
    def is_valid_file(self, filename: str) -> bool:
        """
        Проверяет, является ли строка валидным именем файла.
        
        Критерии:
        - Минимум 3 символа
        - Не входит в blacklist
        - Не является общим словом
        """
        # Проверка длины
        if len(filename) < 3:
            return False
        
        # Проверка blacklist (точное совпадение)
        if filename in self.FILE_BLACKLIST:
            return False
        
        # Проверка вхождения blacklist слов
        for blacklisted in self.FILE_BLACKLIST:
            if blacklisted.lower() in filename.lower() and len(blacklisted) > 5:
                return False
        
        # Дополнительно: проверка что это не просто расширение
        if re.match(r'^\w{1,4}$', filename, re.IGNORECASE):
            return False
        
        return True
    
    def find_all_raw_matches(self, text: str) -> Dict[str, List[Tuple[str, str]]]:
        """
        Находит все "сырые" совпадения для всех типов IOC.
        ВАЖНО: Порядок обработки - сначала URI, потом хеши по убыванию длины.
        
        Returns:
            Словарь {имя_ioc: [(original, cleaned), ...]}
        """
        raw_matches = {}
        
        # Предобработка текста для URI с переносами строк
        processed_text = re.sub(r'([:/])\s*\n\s*', r'\1', text)
        
        # ШАГ 1: Извлекаем URI первыми
        uri_components = {'ips': set(), 'domains': set(), 'files': set()}
        
        for ioc in self.ioc_config:
            if ioc['name'] == 'URI' and ioc.get('enabled', False):
                regex = ioc['regex']
                pattern = re.compile(regex, re.IGNORECASE | re.MULTILINE)
                matches = pattern.findall(processed_text)
                
                uri_pairs = []
                seen = set()
                
                for match in matches:
                    if isinstance(match, tuple):
                        match = match[0]
                    match = match.strip()
                    match = re.sub(r'[.,;!?)}\]]+$', '', match)
                    
                    if match and match not in seen:
                        original = match
                        cleaned = self.clean_ioc(match, 'URI')
                        uri_pairs.append((original, cleaned))
                        seen.add(match)
                        
                        # Извлекаем компоненты для фильтрации
                        components = self.extract_components_from_uri(original)
                        uri_components['ips'].update(components['ips'])
                        uri_components['domains'].update(components['domains'])
                        uri_components['files'].update(components['files'])
                
                raw_matches['URI'] = uri_pairs
                break
        
        # ШАГ 2: Обрабатываем хеши в правильном порядке (SHA256 → SHA1 → MD5)
        hash_order = ['SHA256', 'SHA1', 'MD5']
        found_hash_strings = set()
        
        for hash_name in hash_order:
            for ioc in self.ioc_config:
                if ioc['name'] == hash_name and ioc.get('enabled', False):
                    regex = ioc['regex']
                    pattern = re.compile(regex, re.IGNORECASE | re.MULTILINE)
                    matches = pattern.findall(processed_text)
                    
                    hash_pairs = []
                    seen = set()
                    
                    for match in matches:
                        if isinstance(match, tuple):
                            match = match[0]
                        match = match.strip()
                        
                        # Проверяем, не является ли этот хеш частью уже найденного более длинного
                        is_substring = False
                        for found_hash in found_hash_strings:
                            if match in found_hash:
                                is_substring = True
                                break
                        
                        if not is_substring and match and match not in seen:
                            original = match
                            cleaned = self.clean_ioc(match, hash_name)
                            hash_pairs.append((original, cleaned))
                            seen.add(match)
                            found_hash_strings.add(match)
                    
                    raw_matches[hash_name] = hash_pairs
                    break
        
        # ШАГ 3: Обрабатываем остальные типы IOC
        for ioc in self.ioc_config:
            if not ioc.get('enabled', False):
                continue
            
            name = ioc['name']
            
            # Пропускаем уже обработанные
            if name in ['URI', 'SHA256', 'SHA1', 'MD5']:
                continue
            
            regex = ioc['regex']
            
            try:
                pattern = re.compile(regex, re.IGNORECASE | re.MULTILINE)
                matches = pattern.findall(processed_text)
                
                if matches:
                    # Обработка групп в regex
                    if matches and isinstance(matches[0], tuple):
                        matches = [m[0] if isinstance(m, tuple) else m for m in matches]
                    
                    # Сохраняем пары (original, cleaned)
                    ioc_pairs = []
                    seen = set()
                    
                    for match in matches:
                        match = match.strip()
                        if match and match not in seen:
                            # Убираем trailing знаки препинания
                            match = re.sub(r'[.,;!?)}\]]+$', '', match)
                            
                            # ФИЛЬТРАЦИЯ ПО ТИПАМ
                            should_skip = False
                            
                            # IP: пропустить если из URI
                            if name == 'IP' and match in uri_components['ips']:
                                should_skip = True
                            
                            # DNS: пропустить если из URI или email
                            if name == 'DNS':
                                if match in uri_components['domains']:
                                    should_skip = True
                                # Пропустить если это часть email
                                if match in self.EMAIL_WHITELIST:
                                    should_skip = True
                                # Пропустить если это имя файла (содержит расширение)
                                if re.search(r'\.(rar|zip|exe|pdf|doc)$', match, re.IGNORECASE):
                                    should_skip = True
                            
                            # File: проверка валидности
                            if name == 'File':
                                if not self.is_valid_file(match):
                                    should_skip = True
                                if match in uri_components['files']:
                                    should_skip = True
                            
                            # Email: whitelist
                            if name == 'Email':
                                cleaned_match = self.clean_ioc(match, name)
                                if cleaned_match in self.EMAIL_WHITELIST or match in self.EMAIL_WHITELIST:
                                    should_skip = True
                            
                            if not should_skip:
                                original = match
                                cleaned = self.clean_ioc(match, name)
                                
                                if original and cleaned:
                                    ioc_pairs.append((original, cleaned))
                                    seen.add(match)
                    
                    raw_matches[name] = ioc_pairs
            
            except Exception as e:
                print(f"Ошибка при обработке regex для {name}: {e}")
                raw_matches[name] = []
        
        return raw_matches
    
    def parse(self, file_paths: List[str]) -> Dict[str, List[Tuple[str, str]]]:
        """
        Основной метод парсинга IOC из файлов.
        
        Returns:
            Словарь {имя_ioc: [(original, cleaned), ...]}
        """
        # Извлечение текста
        combined_text = self.extract_from_files(file_paths)
        
        # Извлечение IOC
        ioc_results = self.find_all_raw_matches(combined_text)
        
        # Инициализация пустых списков для неактивных типов
        for ioc in self.ioc_config:
            name = ioc['name']
            if name not in ioc_results:
                ioc_results[name] = []
        
        return ioc_results
