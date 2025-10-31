"""
Модуль для извлечения IOC из документов V2.
Улучшенная обработка [...] и переносов строк.
"""

import re
from typing import List, Dict, Set, Any, Tuple
from docx import Document


class IOCParser:
    """Парсер для извлечения IOC из .docx файлов V2."""
    
    def __init__(self, ioc_config: List[Dict[str, Any]]):
        """Инициализация парсера."""
        self.ioc_config = ioc_config
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлекает весь текст из .docx файла (параграфы и таблицы)."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Извлечение текста из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка при чтении файла {file_path}: {str(e)}")
    
    def extract_from_files(self, file_paths: List[str]) -> str:
        """Извлекает и объединяет текст из нескольких файлов."""
        combined_text = []
        
        for file_path in file_paths:
            text = self.extract_text_from_docx(file_path)
            combined_text.append(text)
        
        return '\n\n'.join(combined_text)
    
    @staticmethod
    def clean_ioc(ioc: str, ioc_type: str) -> str:
        """
        Очищает IOC от обфускации.
        
        Args:
            ioc: Исходный IOC
            ioc_type: Тип IOC
            
        Returns:
            Очищенный IOC
        """
        cleaned = ioc
        
        # Очистка [...] и [:]
        cleaned = cleaned.replace('[.]', '.')
        cleaned = cleaned.replace('[:]', ':')
        cleaned = cleaned.replace('[', '')
        cleaned = cleaned.replace(']', '')
        
        # Дополнительная очистка для URI
        if ioc_type == 'URI':
            # Убираем hxxp -> http, hxxps -> https
            cleaned = cleaned.replace('hxxp://', 'http://')
            cleaned = cleaned.replace('hxxps://', 'https://')
            cleaned = cleaned.replace('hxxp:', 'http:')
            cleaned = cleaned.replace('hxxps:', 'https:')
        
        return cleaned
    
    def find_all_raw_matches(self, text: str) -> Dict[str, List[Tuple[str, str]]]:
        """
        Находит все "сырые" совпадения для всех типов IOC.
        
        Returns:
            Словарь {имя_ioc: [(original, cleaned), ...]}
        """
        raw_matches = {}
        
        # Предобработка текста для URI с переносами строк
        processed_text = re.sub(r'([:/])\s*\n\s*', r'\1', text)
        
        for ioc in self.ioc_config:
            if not ioc.get('enabled', False):
                continue
            
            name = ioc['name']
            regex = ioc['regex']
            
            try:
                pattern = re.compile(regex, re.IGNORECASE | re.MULTILINE)
                matches = pattern.findall(processed_text)
                
                if matches:
                    # Обработка групп в regex
                    if matches and isinstance(matches[0], tuple):
                        matches = [m[0] if isinstance(m, tuple) else m for m in matches]
                    
                    # Сохраняем пары (original, cleaned)
                    ioc_pairs = []
                    seen = set()
                    
                    for match in matches:
                        match = match.strip()
                        if match and match not in seen:
                            # Убираем trailing знаки препинания
                            match = re.sub(r'[.,;!?)}\]]+$', '', match)
                            
                            original = match
                            cleaned = self.clean_ioc(match, name)
                            
                            if original and cleaned:
                                ioc_pairs.append((original, cleaned))
                                seen.add(match)
                    
                    raw_matches[name] = ioc_pairs
            
            except Exception as e:
                print(f"Ошибка при обработке regex для {name}: {e}")
                raw_matches[name] = []
        
        return raw_matches
    
    def filter_overlapping_iocs(self, raw_matches: Dict[str, List[Tuple[str, str]]]) -> Dict[str, List[Tuple[str, str]]]:
        """Фильтрует перекрывающиеся IOC."""
        filtered = {}
        
        # Создаем копии для работы
        for key, values in raw_matches.items():
            filtered[key] = values.copy()
        
        # Фильтр 1: Удаляем IP, которые являются частью URI
        if 'IP' in filtered and 'URI' in filtered:
            uris_cleaned = [cleaned for _, cleaned in filtered['URI']]
            ips_to_remove = []
            
            for idx, (ip_orig, ip_clean) in enumerate(filtered['IP']):
                for uri in uris_cleaned:
                    if ip_clean in uri:
                        ips_to_remove.append(idx)
                        break
            
            filtered['IP'] = [pair for idx, pair in enumerate(filtered['IP']) if idx not in ips_to_remove]
        
        # Фильтр 2: Удаляем DNS, которые являются частью URI
        if 'DNS' in filtered and 'URI' in filtered:
            uris_cleaned = [cleaned for _, cleaned in filtered['URI']]
            dns_to_remove = []
            
            for idx, (dns_orig, dns_clean) in enumerate(filtered['DNS']):
                for uri in uris_cleaned:
                    if dns_clean in uri:
                        dns_to_remove.append(idx)
                        break
            
            filtered['DNS'] = [pair for idx, pair in enumerate(filtered['DNS']) if idx not in dns_to_remove]
        
        # Фильтр 3: Удаляем DNS, которые на самом деле являются именами файлов
        if 'DNS' in filtered and 'File' in filtered:
            files_cleaned = [cleaned for _, cleaned in filtered['File']]
            dns_to_remove = []
            
            for idx, (dns_orig, dns_clean) in enumerate(filtered['DNS']):
                for file in files_cleaned:
                    if dns_clean == file or dns_clean in file:
                        dns_to_remove.append(idx)
                        break
            
            filtered['DNS'] = [pair for idx, pair in enumerate(filtered['DNS']) if idx not in dns_to_remove]
        
        # Фильтр 4: Удаляем MD5, которые являются частью SHA256/SHA1
        if 'MD5' in filtered:
            longer_hashes = []
            if 'SHA256' in filtered:
                longer_hashes.extend([cleaned for _, cleaned in filtered['SHA256']])
            if 'SHA1' in filtered:
                longer_hashes.extend([cleaned for _, cleaned in filtered['SHA1']])
            
            md5_to_remove = []
            for idx, (md5_orig, md5_clean) in enumerate(filtered['MD5']):
                for long_hash in longer_hashes:
                    if md5_clean.lower() in long_hash.lower():
                        md5_to_remove.append(idx)
                        break
            
            filtered['MD5'] = [pair for idx, pair in enumerate(filtered['MD5']) if idx not in md5_to_remove]
        
        return filtered
    
    def parse(self, file_paths: List[str]) -> Dict[str, List[Tuple[str, str]]]:
        """
        Главный метод для парсинга IOC из файлов.
        
        Returns:
            Словарь {имя_ioc: [(original, cleaned), ...]} в порядке приоритета
        """
        # Извлекаем текст из всех файлов
        combined_text = self.extract_from_files(file_paths)
        
        # Находим все совпадения
        raw_matches = self.find_all_raw_matches(combined_text)
        
        # Фильтруем перекрытия
        filtered_matches = self.filter_overlapping_iocs(raw_matches)
        
        # Сортируем в порядке из конфига
        result = {}
        for ioc in self.ioc_config:
            if not ioc.get('enabled', False):
                continue
            
            name = ioc['name']
            if name in filtered_matches and filtered_matches[name]:
                # Сортируем пары для консистентности
                result[name] = sorted(filtered_matches[name], key=lambda x: x[1])
        
        return result
