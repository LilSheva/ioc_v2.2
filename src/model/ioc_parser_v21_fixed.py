"""Парсер IOC из документов Word (.docx) с поддержкой режимов ФСТЕК и ГосСОПКА."""

import os
import re
from typing import List, Dict, Any, Tuple
from docx import Document


class IOCParser:
    """Парсер для извлечения IOC из .docx файлов V2.3 FINAL."""

    def __init__(self, ioc_config: List[Dict[str, Any]], mode: str = "fstek"):
        self.ioc_config = ioc_config
        self.mode = mode
        self.file_blacklist = []
        self.filename_exclusions = []
        for ioc in self.ioc_config:
            if ioc['name'] == 'File':
                self.file_blacklist = ioc.get('file_blacklist', [])
                self.filename_exclusions = ioc.get('filename_exclusions', [])
                break

    def extract_metadata_from_docx(self, file_path: str) -> dict:
        """Извлекает метаданные из начала .docx файла."""
        metadata = {
            "filename": os.path.basename(file_path),
            "bulletin_num": "",
            "event_type": ""
        }

        try:
            doc = Document(file_path)
            header_text = []
            for i, paragraph in enumerate(doc.paragraphs):
                if i < 20:
                    header_text.append(paragraph.text)
                else:
                    break

            full_header = '\n'.join(header_text)

            num_match = re.search(r'№\s*([^\n]+)', full_header)
            if num_match:
                metadata["bulletin_num"] = num_match.group(1).strip()

            event_match = re.search(r'Тип\s+события:\s*(.+?)(?=Источник\s+информации|$)', full_header, re.IGNORECASE | re.DOTALL)
            if event_match:
                event_type = event_match.group(1).strip()
                event_type = re.sub(r'\s+', ' ', event_type)
                metadata["event_type"] = event_type

        except Exception:
            pass

        return metadata

    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлекает весь текст из .docx файла (параграфы и таблицы)."""
        try:
            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    seen_tcs = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id not in seen_tcs and cell.text.strip():
                            text_parts.append(cell.text.strip() + '\n')
                            seen_tcs.add(tc_id)

            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка при чтении файла {file_path}: {str(e)}")

    def extract_from_files(self, file_paths: List[str]) -> str:
        """Извлекает и объединяет текст из нескольких файлов."""
        combined_text = [self.extract_text_from_docx(fp) for fp in file_paths]
        return '\n\n'.join(combined_text)

    @staticmethod
    def clean_ioc(ioc: str, ioc_type: str) -> str:
        """Очищает IOC от обфускации."""
        cleaned = ioc.strip()

        if ioc_type == 'File':
            cleaned = re.sub(r'\s+', ' ', cleaned)

        cleaned = cleaned.replace('[.]', '.')
        cleaned = cleaned.replace('[:]', ':')
        cleaned = cleaned.replace('[', '').replace(']', '')

        if ioc_type == 'URI':
            if '://' in cleaned:
                cleaned = 'http' + cleaned[cleaned.find('://'):]

        return cleaned

    def _get_ioc_filters(self, ioc_name: str) -> Tuple[List[str], List[str]]:
        """Возвращает (blacklist, exclusions) для типа IOC. Значения lower-cased."""
        for ioc in self.ioc_config:
            if ioc['name'] == ioc_name:
                bl = [v.lower() for v in ioc.get('blacklist', [])]
                excl = [v.lower() for v in ioc.get('exclusions', [])]
                return bl, excl
        return [], []

    def _passes_filters(self, cleaned: str, match_start: int, working_text: str,
                        blacklist: List[str], exclusions: List[str]) -> bool:
        """Возвращает True если IOC проходит фильтры blacklist и exclusions."""
        if blacklist and cleaned.lower() in blacklist:
            return False
        if exclusions:
            text_before = working_text[max(0, match_start - 30):match_start].lower().rstrip()
            if any(text_before.endswith(exc) for exc in exclusions):
                return False
        return True

    def _detect_ioc_status(self, full_text: str, ioc_original: str) -> str:
        """Определяет статус IOC по контексту (только для ГосСОПКА)."""
        pos = full_text.find(ioc_original)
        if pos == -1:
            return "block"

        after_text = full_text[pos + len(ioc_original):pos + len(ioc_original) + 150]
        after_lines = after_text.strip().split('\n')
        for line in after_lines[:2]:
            if re.search(r'разблокиров', line, re.IGNORECASE):
                return "unblock"

        before_text = full_text[max(0, pos - 300):pos].lower()

        if 'легитимный' in before_text:
            return "unblock"

        if 'для поиска и блокировки' in before_text:
            return "block"

        if 'для поиска' in before_text:
            return "search"

        return "block"

    def extract_bdu_identifiers(self, file_paths: List[str]) -> List[Tuple[str, str]]:
        """Извлекает BDU-идентификаторы из файлов. Возвращает [(bdu_id, filename)]."""
        bdu_list = []
        bdu_regex = re.compile(r'BDU:\d{4}-\d{4,6}')
        for fp in file_paths:
            text = self.extract_text_from_docx(fp)
            matches = bdu_regex.findall(text)
            fname = os.path.basename(fp)
            for m in set(matches):
                bdu_list.append((m, fname))
        return sorted(set(bdu_list))

    def _extract_with_finditer(self, working_text: str, pattern: str, ioc_name: str) -> Tuple[List[Tuple[str, str]], str]:
        """
        Универсальный экстрактор через finditer с поддержкой blacklist/exclusions.
        Обрабатывает матчи справа налево для корректной позиционной замены.
        Возвращает (pairs, updated_working_text).
        """
        blacklist, exclusions = self._get_ioc_filters(ioc_name)
        seen: set = set()
        collected: List[Tuple[str, str]] = []

        matches = list(re.finditer(pattern, working_text))
        for m in reversed(matches):
            original = m.group(0)
            cleaned = self.clean_ioc(original, ioc_name)
            start, end = m.span()

            should_include = (
                cleaned not in seen
                and self._passes_filters(cleaned, start, working_text, blacklist, exclusions)
            )

            # Заменяем на пробелы в любом случае — предотвращаем повторное извлечение
            working_text = working_text[:start] + ' ' * (end - start) + working_text[end:]

            if should_include:
                seen.add(cleaned)
                collected.append((original, cleaned))

        collected.reverse()
        return collected, working_text

    def find_all_raw_matches(self, text: str) -> Dict[str, List[Tuple[str, str]]]:
        """
        Извлекает IOC из текста в порядке: Email -> URI -> IP -> Files -> DNS -> Hashes.
        """
        raw_matches = {}
        working_text = text

        # ── Email ──────────────────────────────────────────────────
        email_pairs = []
        for ioc in self.ioc_config:
            if ioc['name'] == 'Email' and ioc.get('enabled', False):
                blacklist, exclusions = self._get_ioc_filters('Email')

                if self.mode == "gossopka":
                    # Сначала обфусцированные
                    # Разрешаем [.] и . в local-part, чтобы поймать abc[.]qwe[.]123@mail[.]ru целиком.
                    # \b ненадёжен на стыке ]→буква, используем явные lookaround по классу символов.
                    email_regex_mixed = (
                        r'(?<![A-Za-z0-9._%+\-\[\]])'
                        r'(?:[a-zA-Z0-9_%+\-]+(?:\[\.\]|\.))*'
                        r'[a-zA-Z0-9_%+\-]+'
                        r'@[a-zA-Z0-9-]+(?:(?:\.|\[\.\])[a-zA-Z0-9-]+)+'
                        r'(?![a-zA-Z0-9\-])'
                    )
                    matches_mixed = re.findall(email_regex_mixed, working_text)
                    # Сортируем по убыванию длины — защита от порчи длинного матча коротким при str.replace
                    for match in sorted(set(matches_mixed), key=len, reverse=True):
                        if '[.]' not in match or match not in working_text:
                            continue
                        cleaned = self.clean_ioc(match, 'Email')
                        pos = working_text.find(match)
                        if not self._passes_filters(cleaned, pos, working_text, blacklist, exclusions):
                            working_text = working_text.replace(match, ' ' * len(match))
                            continue
                        email_pairs.append((match, cleaned))
                        working_text = working_text.replace(match, ' ' * len(match))

                    # Затем необфусцированные
                    email_regex_plain = r'\b[a-zA-Z0-9._%+-]+@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}\b'
                    matches_plain = re.findall(email_regex_plain, working_text)
                    for match in sorted(set(matches_plain), key=len, reverse=True):
                        if match not in working_text:
                            continue
                        cleaned = match
                        pos = working_text.find(match)
                        if not self._passes_filters(cleaned, pos, working_text, blacklist, exclusions):
                            working_text = working_text.replace(match, ' ' * len(match))
                            continue
                        email_pairs.append((match, cleaned))
                        working_text = working_text.replace(match, ' ' * len(match))
                else:
                    # ФСТЕК — только обфусцированные.
                    # Разрешаем [.] и . в local-part (домен обязан содержать [.]).
                    email_regex = (
                        r'(?<![A-Za-z0-9._%+\-\[\]])'
                        r'(?:[a-zA-Z0-9_%+\-]+(?:\[\.\]|\.))*'
                        r'[a-zA-Z0-9_%+\-]+'
                        r'@(?:[a-zA-Z0-9-]+\.)*[a-zA-Z0-9-]+\[\.\][a-zA-Z]{2,}'
                        r'(?![a-zA-Z0-9\-])'
                    )
                    matches = re.findall(email_regex, working_text)
                    for match in sorted(set(matches), key=len, reverse=True):
                        if match not in working_text:
                            continue
                        cleaned = self.clean_ioc(match, 'Email')
                        pos = working_text.find(match)
                        if not self._passes_filters(cleaned, pos, working_text, blacklist, exclusions):
                            working_text = working_text.replace(match, ' ' * len(match))
                            continue
                        email_pairs.append((match, cleaned))
                        working_text = working_text.replace(match, ' ' * len(match))
                break
        raw_matches['Email'] = email_pairs

        # ── URI ────────────────────────────────────────────────────
        # Склейка разорванных переносом строки
        stitching_pattern = re.compile(r'([/\]:.)])[ \t]*\n[ \t]*(?=[a-z0-9/\[(])')
        working_text = stitching_pattern.sub(r'\1', working_text)

        uri_pairs = []
        if self.mode == "gossopka":
            uri_pattern = re.compile(r'\b[a-zA-Z0-9][^\s<>"\n\r]*?(?:\[:\]|:)//(?:[^\s<>"\n\r]|[\[].{1,2}[\]])+(?<![.,;])')
        else:
            uri_pattern = re.compile(r'\b[a-zA-Z0-9][^\s<>"\n\r]*?\[:\]//(?:[^.,;\s<>\[\]\n\r]|[\[].{1,2}[\]])+')

        blacklist_uri, exclusions_uri = self._get_ioc_filters('URI')
        uri_matches = sorted(uri_pattern.finditer(working_text), key=lambda m: m.start(), reverse=True)
        for match in uri_matches:
            original = match.group(0)
            cleaned = self.clean_ioc(original, 'URI')
            start, end = match.span()
            if self._passes_filters(cleaned, start, working_text, blacklist_uri, exclusions_uri):
                uri_pairs.append((original, cleaned))
            working_text = working_text[:start] + ' ' * (end - start) + working_text[end:]

        raw_matches['URI'] = sorted(uri_pairs, key=lambda x: x[0])

        # ── IP ─────────────────────────────────────────────────────
        ip_pairs = []
        for ioc in self.ioc_config:
            if ioc['name'] == 'IP' and ioc.get('enabled', False):
                pairs, working_text = self._extract_with_finditer(working_text, ioc['regex'], 'IP')
                ip_pairs.extend(pairs)
                break
        raw_matches['IP'] = ip_pairs

        # ── Files ──────────────────────────────────────────────────
        file_pairs = []
        for match_obj in re.finditer(r'«([^»]+)»', working_text):
            filename = match_obj.group(1)

            start_pos = match_obj.start()
            text_before = working_text[max(0, start_pos - 20):start_pos].lower().rstrip()
            if any(text_before.endswith(word) for word in self.file_blacklist):
                continue

            if filename.strip() in self.filename_exclusions:
                continue

            parts = filename.rsplit('.', 1)
            if len(parts) == 2 and parts[0].strip() and re.match(r'^[a-zA-Z]+$', parts[1].strip()):
                cleaned = self.clean_ioc(filename, 'File')
                file_pairs.append((filename, cleaned))

        for original, _ in file_pairs:
            working_text = working_text.replace(f'«{original}»', ' ' * (len(original) + 2))

        raw_matches['File'] = file_pairs

        # ── DNS ────────────────────────────────────────────────────
        dns_pairs = []
        for ioc in self.ioc_config:
            if ioc['name'] == 'DNS' and ioc.get('enabled', False):
                if self.mode == "gossopka":
                    dns_regex = r'\b(?=[^\s]*\[\.\])[a-zA-Z0-9-]+(?:(?:\.|\[\.\])[a-zA-Z0-9-]+)+\b'
                else:
                    dns_regex = r'\b[a-zA-Z0-9-]+(?:\[\.\][a-zA-Z0-9-]+)+\b'

                blacklist_dns, exclusions_dns = self._get_ioc_filters('DNS')
                matches_dns = list(re.finditer(dns_regex, working_text))
                for m in reversed(matches_dns):
                    original = m.group(0)
                    if '@' in original:
                        continue
                    cleaned = self.clean_ioc(original, 'DNS')
                    start, end = m.span()
                    if self._passes_filters(cleaned, start, working_text, blacklist_dns, exclusions_dns):
                        dns_pairs.append((original, cleaned))
                    working_text = working_text[:start] + ' ' * (end - start) + working_text[end:]
                dns_pairs.reverse()
                break
        raw_matches['DNS'] = dns_pairs

        # ── Hashes ─────────────────────────────────────────────────
        hash_order = ['SHA256', 'SHA1', 'MD5']
        for hash_name in hash_order:
            for ioc in self.ioc_config:
                if ioc['name'] == hash_name and ioc.get('enabled', False):
                    pairs, working_text = self._extract_with_finditer(working_text, ioc['regex'], hash_name)
                    if hash_name not in raw_matches:
                        raw_matches[hash_name] = []
                    raw_matches[hash_name].extend(pairs)
                    break

        return raw_matches

    def parse(self, file_paths: List[str]) -> Dict[str, List[Tuple[str, str, dict]]]:
        """
        Основной метод парсинга IOC из файлов.
        Возвращает Dict[str, List[Tuple[original, cleaned, metadata]]].
        """
        ioc_results = {}
        for ioc in self.ioc_config:
            ioc_results[ioc['name']] = []

        for file_path in file_paths:
            metadata = self.extract_metadata_from_docx(file_path)
            text = self.extract_text_from_docx(file_path)
            file_ioc_results = self.find_all_raw_matches(text)

            for ioc_type, ioc_list in file_ioc_results.items():
                if ioc_type in ioc_results:
                    for original, cleaned in ioc_list:
                        meta = metadata.copy()
                        if self.mode == "gossopka":
                            meta["status"] = self._detect_ioc_status(text, original)
                        else:
                            meta["status"] = "block"
                        ioc_results[ioc_type].append((original, cleaned, meta))

        return ioc_results
