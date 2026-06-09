"""
Адаптер чтения документов docx (python-docx).
"""

from docx import Document
from ioc_analyzer.ports.document_port import DocumentPort, ParagraphData


class DocxAdapter(DocumentPort):
    """
    Реализация порта DocumentPort для файлов .docx.
    """

    def read_paragraphs(self, file_path: str) -> list[ParagraphData]:
        """
        Считывает абзацы из docx, детектируя горизонтальные границы абзаца.
        """
        try:
            doc = Document(file_path)
            res = []
            for p in doc.paragraphs:
                has_border = False
                xml_str = ""
                if p._p.pPr is not None:
                    xml_str = p._p.pPr.xml
                    if "w:pBdr" in xml_str:
                        has_border = True
                res.append(ParagraphData(text=p.text, has_border=has_border, xml_str=xml_str))
            return res
        except Exception as e:
            raise Exception(f"Ошибка чтения абзацев из docx {file_path}: {e}")

    def read_full_text(self, file_path: str) -> str:
        """
        Считывает весь текст из документа (параграфы + таблицы).
        """
        try:
            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    seen_tcs = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id not in seen_tcs and cell.text.strip():
                            text_parts.append(cell.text.strip() + '\n')
                            seen_tcs.add(tc_id)

            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из docx {file_path}: {e}")
